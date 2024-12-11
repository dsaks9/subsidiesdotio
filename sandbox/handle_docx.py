# from docx import Document
# import json
# from pathlib import Path

# def read_docx(file_path):
#     """
#     Read raw content from a DOCX file
#     """
#     doc = Document(file_path)
#     content = []
    
#     for paragraph in doc.paragraphs:
#         if paragraph.text.strip():  # Only include non-empty paragraphs
#             content.append(paragraph.text)
            
#     return content

# def main():
#     # Define paths
#     docx_path = Path('/Users/delonsaks/Documents/subsidies-dot-io/data/vindsubsidies_2_subsidies.docx')
#     output_path = Path('/Users/delonsaks/Documents/subsidies-dot-io/data/vindsubsidies_2_raw.json')
    
#     # Create output directory if it doesn't exist
#     output_path.parent.mkdir(parents=True, exist_ok=True)
    
#     try:
#         # Read the DOCX file
#         content = read_docx(docx_path)
        
#         # Save raw content to JSON file
#         with open(output_path, 'w', encoding='utf-8') as f:
#             json.dump(content, f, ensure_ascii=False, indent=2)
            
#         print(f"Successfully saved raw content with {len(content)} paragraphs to {output_path}")
        
#     except Exception as e:
#         print(f"Error processing file: {e}")

# if __name__ == "__main__":
#     main()

from docx import Document
import re
from pathlib import Path
def extract_subsidy_information(docx_path):
    # Load the .docx file
    doc = Document(docx_path)
    
    # Initialize variables
    subsidies = []
    current_subsidy = {}
    current_key = None
    
    # Helper function to detect new subsidy names
    def is_new_subsidy(line):
        # Assuming subsidy names are all uppercase or bold
        return bool(re.match(r"^[A-Z ]+$", line.strip())) or len(line.split()) < 5
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue  # Skip empty lines
        
        # Check for new subsidy start
        if is_new_subsidy(text):
            # Save the previous subsidy if we have one
            if current_subsidy:
                subsidies.append(current_subsidy)
                current_subsidy = {}
            
            # Start a new subsidy entry
            current_subsidy["Subsidy Name"] = text
            current_key = None  # Reset current key
        else:
            # Check if this is a new key
            match = re.match(r"^(.*?):$", text)
            if match:
                # Save the previous key-value pair if applicable
                if current_key and current_key in current_subsidy:
                    current_subsidy[current_key] = current_subsidy[current_key].strip()
                
                # Start a new key
                current_key = match.group(1)
                current_subsidy[current_key] = ""
            elif current_key:
                # Append to the current key
                current_subsidy[current_key] += text + " "
    
    # Add the last subsidy to the list
    if current_subsidy:
        subsidies.append(current_subsidy)
    
    return subsidies

# Example usage
docx_file_path = "/Users/delonsaks/Documents/subsidies-dot-io/data/vindsubsidies_2_subsidies.docx"
subsidy_data = extract_subsidy_information(docx_file_path)

output_path = Path('/Users/delonsaks/Documents/subsidies-dot-io/data/vindsubsidies_2_parsed.json')


# Display the structured data
import json

with open(output_path, 'w', encoding='utf-8') as f:
    json.dump(subsidy_data, f, ensure_ascii=False, indent=2)
# print(json.dumps(subsidy_data, indent=2, ensure_ascii=False))
